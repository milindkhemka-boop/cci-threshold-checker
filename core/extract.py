"""
Extract candidate financial figures from uploaded documents.

Supported in-memory: .pdf (pypdf), .xlsx (openpyxl), .docx (python-docx),
.eml/.txt/.csv/.html (standard library). Optional parsers degrade gracefully —
if a library is missing the user can still paste text or type figures in.

Extraction is a *human-in-the-loop assist*: it surfaces likely numbers with the
line they came from, and the user confirms / assigns them to the right field.
Nothing is stored — bytes are parsed and discarded.
"""

import io
import re
from email import message_from_bytes
from html.parser import HTMLParser

from . import numbers as N

# Keyword -> the metric we think the line represents. Order = display priority.
METRIC_KEYWORDS = [
    ("turnover", [r"revenue from operations", r"\bturnover\b", r"revenue from operation"]),
    ("total_revenue", [r"total revenue", r"total income", r"revenue from operations and other income"]),
    ("total_assets", [r"total assets", r"total non-current and current assets"]),
    ("other_income", [r"other income"]),
    ("profit", [r"profit before tax", r"profit for the (year|period)", r"net profit"]),
    ("net_worth", [r"net worth", r"total equity", r"shareholders[' ]+funds"]),
    ("deal_value", [r"enterprise value", r"transaction value", r"purchase consideration",
                    r"deal value", r"consideration"]),
]

INDIA_HINT = re.compile(r"\bindia(n)?\b|\bdomestic\b|\bstandalone\b", re.I)
GLOBAL_HINT = re.compile(r"\bglobal\b|\bworldwide\b|\bconsolidat", re.I)

# A money expression: optional currency, the number, optional scale word — so that
# "USD 4 billion" and "₹4,000 crore" each keep their own currency and magnitude.
AMOUNT_RE = re.compile(
    r"(?:₹|rs\.?|inr|us\$|usd|\$|eur|€|gbp|£|jpy|¥|aed|dhs?)?\s*"
    r"\(?-?\d[\d,]*\.?\d*\)?"
    r"\s*(?:lakh\s*crores?|crores?|cr|lakhs?|lacs?|arabs?|trillions?|tn|"
    r"billions?|bn|millions?|mln|mn|thousand|k)?",
    re.I,
)


class _HTMLText(HTMLParser):
    def __init__(self):
        super().__init__()
        self.parts = []

    def handle_data(self, data):
        self.parts.append(data)

    def text(self):
        return " ".join(self.parts)


def _html_to_text(html):
    p = _HTMLText()
    try:
        p.feed(html)
    except Exception:
        return re.sub(r"<[^>]+>", " ", html)
    return p.text()


def extract_text(filename, data):
    """Return (text, note) from raw bytes; note flags any missing-parser fallback."""
    name = (filename or "").lower()
    note = ""
    try:
        if name.endswith(".pdf"):
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(data))
                return "\n".join((pg.extract_text() or "") for pg in reader.pages), note
            except ImportError:
                return "", "PDF parser (pypdf) not installed — paste the text or enter figures manually."
        if name.endswith(".xlsx"):
            try:
                from openpyxl import load_workbook
                wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
                lines = []
                for ws in wb.worksheets:
                    for row in ws.iter_rows(values_only=True):
                        cells = ["" if c is None else str(c) for c in row]
                        if any(cells):
                            lines.append("\t".join(cells))
                return "\n".join(lines), note
            except ImportError:
                return "", "Excel parser (openpyxl) not installed — export to CSV or enter figures manually."
        if name.endswith(".docx"):
            try:
                import docx
                doc = docx.Document(io.BytesIO(data))
                lines = [p.text for p in doc.paragraphs]
                for tbl in doc.tables:
                    for row in tbl.rows:
                        lines.append("\t".join(c.text for c in row.cells))
                return "\n".join(lines), note
            except ImportError:
                return "", "Word parser (python-docx) not installed — paste text or enter figures manually."
        if name.endswith(".eml"):
            msg = message_from_bytes(data)
            chunks = []
            for part in msg.walk():
                ctype = part.get_content_type()
                if ctype == "text/plain":
                    chunks.append(part.get_payload(decode=True).decode("utf-8", "ignore"))
                elif ctype == "text/html":
                    chunks.append(_html_to_text(part.get_payload(decode=True).decode("utf-8", "ignore")))
            if not chunks and not msg.is_multipart():
                chunks.append(data.decode("utf-8", "ignore"))
            return "\n".join(chunks), note
        if name.endswith((".htm", ".html")):
            return _html_to_text(data.decode("utf-8", "ignore")), note
        # .txt, .csv, anything else: best-effort decode
        return data.decode("utf-8", "ignore"), note
    except Exception as exc:
        return "", f"Could not parse {filename}: {exc}"


def find_candidates(text, max_per_metric=2):
    """
    Scan text line-by-line for known metrics and return candidate figures.
    Each candidate: {metric, label, raw, value_inr, currency, scope_hint, line}.
    """
    if not text:
        return [], (1.0, "absolute"), "INR"

    scale_mult, scale_label = N.detect_statement_scale(text)
    currency = N.detect_currency(text)

    lines = [ln.strip() for ln in re.split(r"[\r\n]+", text) if ln.strip()]
    candidates = []
    counts = {}

    for ln in lines:
        low = ln.lower()
        consumed = []  # (start,end) money spans already claimed on this line
        for metric, patterns in METRIC_KEYWORDS:
            if counts.get(metric, 0) >= max_per_metric:
                continue
            # Every position where this metric's keyword appears — so a line like
            # "worldwide turnover USD 4 bn, India turnover ₹4,000 cr" yields both.
            kpositions = sorted({mm.end() for p in patterns for mm in re.finditer(p, low)})
            for kpos in kpositions:
                if counts.get(metric, 0) >= max_per_metric:
                    break
                # First unclaimed money expression at/after this keyword. Each
                # figure keeps its own currency + scale word; % tokens skipped.
                picked = None
                for m in AMOUNT_RE.finditer(ln):
                    if m.end() <= kpos or (m.start(), m.end()) in consumed:
                        continue
                    seg = m.group(0).strip()
                    if not re.search(r"\d", seg) or ln[m.end():m.end() + 1] == "%":
                        continue
                    local_ccy = N.detect_currency(seg, default=currency)
                    val = N.parse_amount(seg, assume_scale=scale_mult)
                    if val is not None and abs(val) >= scale_mult:  # ≥ 1 unit of the scale
                        picked = (seg, val, local_ccy, m.start(), m.end())
                        break
                if picked is None:
                    continue
                seg_text, val, local_ccy, pos, end = picked
                consumed.append((pos, end))
                # Scope from the comma-clause the figure sits in (avoids bleed
                # across clauses, e.g. "...USD 4 bn, India turnover ...").
                left = ln.rfind(",", 0, pos)
                right = ln.find(",", pos)
                clause = ln[(left + 1 if left >= 0 else 0):(right if right >= 0 else len(ln))]
                scope = "india" if INDIA_HINT.search(clause) else ("global" if GLOBAL_HINT.search(clause) else "")
                candidates.append({
                    "metric": metric,
                    "label": _metric_label(metric),
                    "raw": seg_text,
                    "value_inr": val if local_ccy == "INR" else None,
                    "value_native": val,
                    "currency": local_ccy,
                    "scope_hint": scope,
                    "line": ln[:160],
                })
                counts[metric] = counts.get(metric, 0) + 1
    return candidates, (scale_mult, scale_label), currency


def _metric_label(metric):
    return {
        "turnover": "Revenue from operations (turnover)",
        "total_revenue": "Total revenue / income",
        "total_assets": "Total assets",
        "other_income": "Other income",
        "profit": "Profit",
        "net_worth": "Net worth / equity",
        "deal_value": "Deal / transaction value",
    }.get(metric, metric)
